from __future__ import annotations
import base64
import subprocess
import tempfile
import uuid
from pathlib import Path
from typing import List, Mapping, Tuple
import re
import xml.etree.ElementTree as ET

from models import orm
from templates import PaperTemplate

try:
    from docx import Document
except ImportError:  # pragma: no cover - optional dependency
    Document = None

try:
    import cairosvg
except ImportError:  # pragma: no cover - optional dependency
    cairosvg = None


class ExportService:
    """
    导出服务：
    - build_latex: 生成可编译的 LaTeX 文本。
    - compile_pdf: 调用本地 pdflatex，如果不可用则返回错误。
    - build_docx: 使用 python-docx 生成 Word。
    """
    def build_latex_from_template(
        self,
        paper: orm.Paper,
        pq_list: List[orm.PaperQuestion],
        question_map: Mapping[str, orm.Question],
        template: PaperTemplate,
        include_answer: bool = True,
        include_explanation: bool = True,
    ) -> Tuple[str, List[Tuple[str, bytes]]]:
        """
        基于模板分块生成 LaTeX（简化版，使用 enumerate + 分块标题）
        """
        header = r"""\documentclass[12pt]{ctexart}
\usepackage[UTF8,fontset=none]{ctex}
\setCJKmainfont{Noto Serif CJK SC}
\setCJKsansfont{Noto Sans CJK SC}
\setCJKmonofont{Noto Sans Mono CJK SC}
\usepackage{amsmath,amssymb}
\usepackage{geometry,graphicx,enumitem,tikz,fancyhdr}
\usepackage[bodytextleadingratio=1.67,restoremathleading=true]{zhlineskip}
\geometry{paperheight=26cm,paperwidth=18.4cm,left=2cm,right=2cm,top=1.5cm,bottom=2cm,headsep=10pt}
\pagestyle{fancy}
\renewcommand{\headrulewidth}{0pt}
\setlength{\parskip}{0.6em}
\setlength{\parindent}{0pt}
\newcommand{\choicefour}[4]{%%
  \begin{tabular}{p{0.45\textwidth}p{0.45\textwidth}}
  \textsf{A}.~#1 & \textsf{B}.~#2\\
  \textsf{C}.~#3 & \textsf{D}.~#4
  \end{tabular}
}
\begin{document}
\begin{center}\Large %s\end{center}
""" % self._escape_latex(paper.title)

        attachments: List[Tuple[str, bytes]] = []
        body_parts: List[str] = []

        sections = self._sections_for_template(template)
        # 建立 order -> PaperQuestion 映射
        pq_by_order = {pq.order: pq for pq in pq_list}

        for section in sections:
            body_parts.append(r"{\bf %s}" % section["title"])
            body_parts.append(
                r"\begin{enumerate}[label=\arabic*.,start=%d,leftmargin=1.5em,itemsep=1em]"
                % section["start"]
            )
            for slot in section["slots"]:
                pq = pq_by_order.get(slot["order"])
                if not pq:
                    # 缺题直接占位
                    body_parts.append(r"\item (%s分) \textit{缺题占位}" % slot["score"])
                    continue
                q = question_map.get(pq.question_id)
                if not q:
                    body_parts.append(r"\item (%s分) \textit{缺题占位}" % slot["score"])
                    continue
                item_lines = []
                item_lines.append(r"\item (%s分) %s" % (pq.score or slot["score"], self._escape_latex(q.question_text)))
                # 选项
                if q.options and len(q.options) == 4 and (q.question_type or "").startswith("choice"):
                    a, b, c, d = [self._strip_option_prefix(opt) for opt in q.options]
                    item_lines.append(r"\choicefour{%s}{%s}{%s}{%s}" % (
                        self._escape_latex(a),
                        self._escape_latex(b),
                        self._escape_latex(c),
                        self._escape_latex(d),
                    ))
                elif q.options:
                    item_lines.append(r"\begin{enumerate}[label=\Alph*. ,leftmargin=1.2em,itemsep=0.2em]")
                    for opt in q.options:
                        item_lines.append(r"\item %s" % self._escape_latex(self._strip_option_prefix(opt)))
                    item_lines.append(r"\end{enumerate}")
                # 图形
                if q.has_geometry and q.geometry_tikz:
                    item_lines.append(self._wrap_diagram_block(q.geometry_tikz))
                elif q.has_geometry and q.geometry_svg:
                    tikz_block = self._svg_to_tikz_block(q.geometry_svg)
                    if tikz_block:
                        item_lines.append(self._wrap_diagram_block(tikz_block))
                    else:
                        # TikZ 失败，fallback 到 PDF（矢量格式）
                        svg_result = self._svg_to_pdf_attachment(q.geometry_svg)
                        if svg_result:
                            fname, data = svg_result
                            attachments.append((fname, data))
                            img = f'\\includegraphics[width=0.48\\textwidth]{{{fname}}}'
                            item_lines.append(self._wrap_diagram_block(img))
                        else:
                            item_lines.append("\n% SVG 转换失败，未插入图形\n")
                # 答案/解析
                if include_answer and q.answer:
                    item_lines.append(r"\textbf{答案：} %s" % self._escape_latex(q.answer))
                if include_explanation and q.explanation:
                    item_lines.append(r"\textbf{解析：} %s" % self._escape_latex(q.explanation))
                body_parts.append("\n".join(item_lines))
            body_parts.append(r"\end{enumerate}")

        footer = r"\end{document}"
        return header + "\n\n".join(body_parts) + "\n" + footer, attachments

    def _sections_for_template(self, template: PaperTemplate):
        """
        根据模板定义分块标题和槽位
        """
        if template.id == "gaokao_new_1":
            sections = []
            slots_sorted = sorted(template.slots, key=lambda s: s.order)
            # 按 order 划分
            sections.append({
                "title": "一、选择题：本大题共 8 小题，每小题 5 分，共 40 分。",
                "start": 1,
                "slots": [{"order": s.order, "score": s.default_score} for s in slots_sorted if 1 <= s.order <= 8],
            })
            sections.append({
                "title": "二、选择题（多选）：本题共 3 小题，每小题 6 分，共 18 分。",
                "start": 9,
                "slots": [{"order": s.order, "score": s.default_score} for s in slots_sorted if 9 <= s.order <= 11],
            })
            sections.append({
                "title": "三、填空题：本大题共 3 小题，每小题 5 分，共 15 分。",
                "start": 12,
                "slots": [{"order": s.order, "score": s.default_score} for s in slots_sorted if 12 <= s.order <= 14],
            })
            sections.append({
                "title": "四、解答题：本大题共 5 小题，共 77 分。",
                "start": 15,
                "slots": [{"order": s.order, "score": s.default_score} for s in slots_sorted if 15 <= s.order <= 19],
            })
            return sections
        # 默认无分块
        slots_sorted = sorted(template.slots, key=lambda s: s.order)
        return [{
            "title": "试卷",
            "start": slots_sorted[0].order if slots_sorted else 1,
            "slots": [{"order": s.order, "score": s.default_score} for s in slots_sorted],
        }]

    def build_answer_latex(
        self,
        paper: orm.Paper,
        pq_list: List[orm.PaperQuestion],
        question_map: Mapping[str, orm.Question],
    ) -> Tuple[str, List[Tuple[str, bytes]]]:
        """
        生成答案卷 LaTeX：
        - 选择题（单选/多选）：表格显示题号+答案字母
        - 填空题：题号+答案值
        - 解答题：题号+完整答案
        """
        header = r"""\documentclass[12pt]{ctexart}
\usepackage[UTF8,fontset=none]{ctex}
\setCJKmainfont{Noto Serif CJK SC}
\setCJKsansfont{Noto Sans CJK SC}
\setCJKmonofont{Noto Sans Mono CJK SC}
\usepackage{unicode-math}
\setmathfont{Latin Modern Math}
\usepackage{amsmath,amssymb}
\usepackage{geometry,graphicx,enumitem,array,booktabs,tikz,fancyhdr}
\usepackage[bodytextleadingratio=1.67,restoremathleading=true]{zhlineskip}
\geometry{paperheight=26cm,paperwidth=18.4cm,left=2cm,right=2cm,top=1.5cm,bottom=2cm,headsep=10pt}
\pagestyle{fancy}
\renewcommand{\headrulewidth}{0pt}
\setlength{\parskip}{0.6em}
\setlength{\parindent}{0pt}
\providecommand{\SetMathEnvironmentSinglespace}[1]{}
\newcommand{\choicefour}[4]{%
  \begin{tabular}{p{0.45\textwidth}p{0.45\textwidth}}
  \textsf{A}.~#1 & \textsf{B}.~#2\\
  \textsf{C}.~#3 & \textsf{D}.~#4
  \end{tabular}
}
\begin{document}
\SetMathEnvironmentSinglespace{1}
\lineskiplimit=5.5pt
\lineskip=7pt
\abovedisplayshortskip=5pt
\belowdisplayshortskip=5pt
\abovedisplayskip=5pt
\begin{center}\Large\textbf{%s — 答案卷}\end{center}
\vspace{1em}
""" % self._escape_latex(paper.title)

        attachments: List[Tuple[str, bytes]] = []
        body_parts: List[str] = []
        
        # 按 order 排序
        pq_sorted = sorted(pq_list, key=lambda x: x.order)
        
        # 分类题目
        choices = []  # 单选 1-8
        multis = []   # 多选 9-11
        fillblanks = []  # 填空 12-14
        solves = []   # 解答 15-19
        
        for pq in pq_sorted:
            q = question_map.get(pq.question_id)
            if not q:
                continue
            order = pq.order
            if 1 <= order <= 8:
                choices.append((order, q))
            elif 9 <= order <= 11:
                multis.append((order, q))
            elif 12 <= order <= 14:
                fillblanks.append((order, q))
            else:
                solves.append((order, q))
        
        def extract_answer_letter(answer_text: str) -> str:
            """从答案文本中提取选项字母（如 A/B/C/D 或 AB/CD）"""
            if not answer_text:
                return ""
            # 尝试提取【答案】后的内容
            import re
            match = re.search(r'【答案】\s*([A-Za-z]+)', answer_text)
            if match:
                return match.group(1).upper()
            # 直接匹配开头的字母
            match = re.match(r'^([A-Za-z]+)', answer_text.strip())
            if match:
                return match.group(1).upper()
            # 返回前20字符作为fallback
            return answer_text.strip()[:20]
        
        def extract_fillblank_answer(answer_text: str) -> str:
            """从答案文本中提取填空答案"""
            if not answer_text:
                return ""
            import re
            match = re.search(r'【答案】\s*(.+?)(?=【|$)', answer_text, re.DOTALL)
            if match:
                return match.group(1).strip()[:50]
            return answer_text.strip()[:50]
        
        # 一、单选题答案表格
        if choices:
            body_parts.append(r"{\bf 一、选择题答案}")
            body_parts.append(r"\begin{center}")
            body_parts.append(r"\begin{tabular}{|" + "c|" * len(choices) + "}")
            body_parts.append(r"\hline")
            body_parts.append(" & ".join([str(o) for o, _ in choices]) + r" \\")
            body_parts.append(r"\hline")
            body_parts.append(" & ".join([extract_answer_letter(q.answer) for _, q in choices]) + r" \\")
            body_parts.append(r"\hline")
            body_parts.append(r"\end{tabular}")
            body_parts.append(r"\end{center}")
            body_parts.append(r"\vspace{1em}")
        
        # 二、多选题答案表格
        if multis:
            body_parts.append(r"{\bf 二、多选题答案}")
            body_parts.append(r"\begin{center}")
            body_parts.append(r"\begin{tabular}{|" + "c|" * len(multis) + "}")
            body_parts.append(r"\hline")
            body_parts.append(" & ".join([str(o) for o, _ in multis]) + r" \\")
            body_parts.append(r"\hline")
            body_parts.append(" & ".join([extract_answer_letter(q.answer) for _, q in multis]) + r" \\")
            body_parts.append(r"\hline")
            body_parts.append(r"\end{tabular}")
            body_parts.append(r"\end{center}")
            body_parts.append(r"\vspace{1em}")
        
        # 三、填空题答案
        if fillblanks:
            body_parts.append(r"{\bf 三、填空题答案}")
            body_parts.append(r"\begin{enumerate}[label=\arabic*.,start=12,leftmargin=1.5em]")
            for order, q in fillblanks:
                ans = extract_fillblank_answer(q.answer)
                body_parts.append(r"\item %s" % self._escape_latex(ans))
            body_parts.append(r"\end{enumerate}")
            body_parts.append(r"\vspace{1em}")
        
        # 四、解答题完整答案
        if solves:
            body_parts.append(r"{\bf 四、解答题答案}")
            body_parts.append(r"\begin{enumerate}[label=\arabic*.,start=15,leftmargin=1.5em,itemsep=1.5em]")
            for order, q in solves:
                body_parts.append(r"\item %s" % self._escape_latex(q.answer or "（无答案）"))
            body_parts.append(r"\end{enumerate}")
        
        footer = r"\end{document}"
        return header + "\n\n".join(body_parts) + "\n" + footer, attachments

    def build_single_question_latex(
        self,
        question: dict,
        include_answer: bool = True,
        include_explanation: bool = True,
    ) -> Tuple[str, List[Tuple[str, bytes]]]:
        """
        为单题生成可编译的 LaTeX 片段，返回 (latex, attachments)
        """
        header = r"""\documentclass[12pt,a4paper]{article}
\usepackage{ctex}
\usepackage{amsmath,amssymb}
\usepackage{geometry}
\usepackage{graphicx}
\usepackage{tikz}
\geometry{left=2cm,right=2cm,top=2.5cm,bottom=2.5cm}
\begin{document}
"""
        body_parts: List[str] = []
        attachments: List[Tuple[str, bytes]] = []

        body_parts.append(self._escape_latex(question.get("questionText") or ""))
        options = question.get("options") or []
        for opt in options:
            body_parts.append(r"\par " + self._escape_latex(opt))

        if question.get("hasGeometry") and question.get("geometryTikz"):
            body_parts.append(self._wrap_diagram_block(question.get("geometryTikz")))
        elif question.get("hasGeometry") and question.get("geometrySvg"):
            # 优先尝试 TikZ，失败则用 PDF
            tikz_block = self._svg_to_tikz_block(question.get("geometrySvg"))
            if tikz_block:
                body_parts.append(self._wrap_diagram_block(tikz_block))
            else:
                svg_result = self._svg_to_pdf_attachment(question.get("geometrySvg"))
                if svg_result:
                    fname, data = svg_result
                    attachments.append((fname, data))
                    img = f'\\includegraphics[width=0.48\\textwidth]{{{fname}}}'
                    body_parts.append(self._wrap_diagram_block(img))
                else:
                    body_parts.append("\n% SVG 转换失败，未插入图形\n")

        if include_answer and question.get("answer"):
            body_parts.append(f"\n\\textbf{{答案：}} {self._escape_latex(question.get('answer') or '')}")
        if include_explanation and question.get("explanation"):
            body_parts.append(f"\n\\textbf{{解析：}} {self._escape_latex(question.get('explanation') or '')}")

        footer = r"""\end{document}
"""
        latex = header + "\n\n".join(body_parts) + footer
        return latex, attachments

    async def export_stub(self, paper_id: str, fmt: str, paper_payload: dict | None = None):
        return {
            "paperId": paper_id,
            "format": fmt,
            "status": "not_implemented",
            "message": "导出服务尚未实现，返回渲染用内容。",
            "paper": paper_payload,
            "latex": paper_payload.get("latex") if paper_payload else None,
        }

    def build_latex(
        self,
        paper: orm.Paper,
        pq_list: List[orm.PaperQuestion],
        question_map: Mapping[str, orm.Question],
        include_answer: bool = True,
        include_explanation: bool = True,
    ) -> Tuple[str, List[Tuple[str, bytes]]]:
        """
        生成高考卷风格的 LaTeX，参照 2025juan1.tex 模板
        """
        # 高考卷标准 LaTeX Header
        header = r"""\documentclass[no-math]{ctexart}
\setCJKmainfont{Noto Serif CJK SC}
\setCJKsansfont{Noto Sans CJK SC}
\setCJKmonofont{Noto Sans Mono CJK SC}
\everymath{\displaystyle}

\usepackage{amsmath,amssymb}
\usepackage{tikz}
\usetikzlibrary{arrows.meta,patterns,calc}
\usepackage{graphicx}
\usepackage{enumitem}
\setenumerate{itemsep=0pt,partopsep=0pt,parsep=\parskip,topsep=0pt}
\allowdisplaybreaks[4]
\tikzset{
  every picture/.style={scale=0.75},
  every node/.style={font=\small},
  line width=0.5pt,
  >={Stealth[length=4pt]}
}

\usepackage[paperheight=26cm,paperwidth=18.4cm,left=2cm,right=2cm,top=1.5cm,bottom=2cm,headsep=10pt]{geometry}
\usepackage{fancyhdr}
\pagestyle{fancy}
\renewcommand{\headrulewidth}{0pt}
\usepackage{lastpage}
\usepackage[bodytextleadingratio=1.67,restoremathleading=true]{zhlineskip}
\usepackage{ifthen}

%% 选项自适应排版命令
\newcommand{\onech}[4]{\makebox[3.4cm][l]{{\sf A}．#1}\makebox[3.4cm][l]{{\sf B}．#2}\makebox[3.4cm][l]{{\sf C}．#3}\makebox[3.4cm][l]{{\sf D}．#4}}
\newcommand{\twoch}[4]{\makebox[6.8cm][l]{{\sf A}．#1}\makebox[6.8cm][l]{{\sf B}．#2}\\ \makebox[6.8cm][l]{{\sf C}．#3}\makebox[6.8cm][l]{{\sf D}．#4}}
\newcommand{\fourch}[4]{{\sf A}．#1\\ {\sf B}．#2\\ {\sf C}．#3\\ {\sf D}．#4}

\newlength\widthcha
\newlength\widthchb
\newlength\widthch
\newlength\fourthtabwidth
\setlength\fourthtabwidth{0.22\textwidth}
\newlength\halftabwidth
\setlength\halftabwidth{0.45\textwidth}

\newcommand{\choice}[4]{%%
  \settowidth\widthcha{{\sf A}M.#1}%%
  \setlength{\widthch}{\widthcha}%%
  \settowidth\widthchb{{\sf B}M.#2}%%
  \ifthenelse{\lengthtest{\widthch < \widthchb}}{\setlength{\widthch}{\widthchb}}{}%%
  \settowidth\widthchb{{\sf C}M.#3}%%
  \ifthenelse{\lengthtest{\widthch < \widthchb}}{\setlength{\widthch}{\widthchb}}{}%%
  \settowidth\widthchb{{\sf D}M.#4}%%
  \ifthenelse{\lengthtest{\widthch < \widthchb}}{\setlength{\widthch}{\widthchb}}{}%%
  \ifthenelse{\lengthtest{\widthch < \fourthtabwidth}}{\onech{#1}{#2}{#3}{#4}}%%
  {\ifthenelse{\lengthtest{\widthch < \halftabwidth}}{\twoch{#1}{#2}{#3}{#4}}%%
  {\fourch{#1}{#2}{#3}{#4}}}%%
}

%% 填空横线（兼容数学模式和文本模式）
\newcommand{\undsp}{\underline{\makebox[3em]{}}}

%% 斜着的平行符号（高考风格）
\newcommand{\spar}{\mathrel{/\mkern-5mu/}}

\begin{document}
\SetMathEnvironmentSinglespace{1}
\lineskiplimit=5.5pt
\lineskip=7pt
\abovedisplayshortskip=5pt
\belowdisplayshortskip=5pt
\abovedisplayskip=5pt
\belowdisplayskip=5pt

\fancyfoot[C]{\bf\sf 数学试题 第{\sf\thepage} 页 （共~{\sf\pageref{LastPage}}~页）}

\begin{center}
\zihao{2}\heiti %s
\end{center}
""" % self._escape_latex(paper.title)

        body_parts = []
        attachments: List[Tuple[str, bytes]] = []
        
        # 按题型分组（高考卷结构：单选、多选、填空、解答）
        SECTION_ORDER = ['choice_single', 'choice_multi', 'fill', 'solve']
        SECTION_INFO = {
            'choice_single': {'name': '一', 'title': '选择题：本题共 %d 小题，每小题 %d 分，共 %d 分。在每小题给出的四个选项中，只有一项是符合题目要求的。'},
            'choice_multi': {'name': '二', 'title': '选择题：本题共 %d 小题，每小题 %d 分，共 %d 分。在每小题给出的选项中，有多项符合题目要求。'},
            'fill': {'name': '三', 'title': '填空题：本题共 %d 小题，每小题 %d 分，共 %d 分。'},
            'solve': {'name': '四', 'title': '解答题：本题共 %d 小题，共 %d 分。解答应写出文字说明、证明过程或演算步骤。'},
        }
        
        # 收集题目按类型分组
        questions_by_type: dict = {}
        for pq in sorted(pq_list, key=lambda x: x.order):
            q = question_map.get(pq.question_id)
            if not q:
                continue
            qtype = q.question_type or 'solve'
            # 归类题型
            if qtype in ('choice', 'choice_single'):
                qtype = 'choice_single'
            elif qtype in ('multi', 'choice_multi'):
                qtype = 'choice_multi'
            elif qtype in ('fillblank', 'fill'):
                qtype = 'fill'
            else:
                qtype = 'solve'
            if qtype not in questions_by_type:
                questions_by_type[qtype] = []
            questions_by_type[qtype].append((pq, q))
        
        # 动态调整章节编号
        section_number = 0
        question_number = 0
        
        for section_type in SECTION_ORDER:
            if section_type not in questions_by_type:
                continue
            
            section_number += 1
            section_questions = questions_by_type[section_type]
            section_count = len(section_questions)
            section_scores = [pq.score for pq, _ in section_questions]
            total_score = sum(section_scores)
            avg_score = section_scores[0] if section_scores else 5
            
            section_content = []
            
            # 章节标题
            info = SECTION_INFO.get(section_type, SECTION_INFO['solve'])
            section_names = ['一', '二', '三', '四', '五']
            sec_name = section_names[section_number - 1] if section_number <= 5 else str(section_number)
            
            if section_type == 'solve':
                title = info['title'] % (section_count, total_score)
            else:
                title = info['title'] % (section_count, avg_score, total_score)
            
            section_content.append(r"\begin{enumerate}[align=left,labelindent=0em,labelwidth=2em,labelsep=0em,leftmargin=2em]")
            section_content.append(r"\item[{\bf %s、}]{\bf\sf %s}" % (sec_name, title))
            section_content.append(r"\end{enumerate}")
            
            # 题目列表（紧凑布局）
            section_content.append(r"\begin{enumerate}[align=left,labelindent=0em,label={\bf\sf\arabic*.},labelwidth=1.5em,labelsep=0em,leftmargin=1.5em,itemsep=0pt,topsep=0pt,start=%d]" % (question_number + 1))
            
            for pq, q in section_questions:
                question_number += 1
                try:
                    item_parts = []
                    
                    # 题干
                    escaped_text = self._escape_latex(q.question_text)
                    item_parts.append(r"\item " + escaped_text)
                    
                    # 选项（选择题）
                    if section_type in ('choice_single', 'choice_multi') and q.options and len(q.options) == 4:
                        a, b, c, d = [self._escape_latex(self._strip_option_prefix(opt)) for opt in q.options]
                        item_parts.append(r"\\" + "\n" + r"\choice{%s}{%s}{%s}{%s}" % (a, b, c, d))
                    elif q.options:
                        # 非标准选项数量
                        item_parts.append(r"\\")
                        for i, opt in enumerate(q.options):
                            label = chr(ord('A') + i)
                            item_parts.append(r"{\sf %s}．%s\quad" % (label, self._escape_latex(self._strip_option_prefix(opt))))
                    
                    # 获取图形内容（如果有）
                    diagram_content = None
                    if q.has_geometry and q.geometry_tikz:
                        diagram_content = q.geometry_tikz
                    elif q.has_geometry and q.geometry_svg:
                        tikz_block = self._svg_to_tikz_block(q.geometry_svg)
                        if tikz_block:
                            diagram_content = tikz_block
                        else:
                            # TikZ 失败，fallback 到 PDF（矢量格式）
                            svg_result = self._svg_to_pdf_attachment(q.geometry_svg)
                            if svg_result:
                                fname, data = svg_result
                                attachments.append((fname, data))
                                diagram_content = f'\\includegraphics[width=0.35\\textwidth]{{{fname}}}'
                    
                    # 根据题型决定图形布局
                    if section_type == 'solve':
                        # 解答题：图在留白左侧（左图右留白）
                        if diagram_content and not include_answer:
                            item_parts.append("\n" + r"\par\noindent")
                            item_parts.append(r"\begin{minipage}[t]{0.45\textwidth}")
                            item_parts.append(r"\centering")
                            item_parts.append(diagram_content)
                            item_parts.append(r"\end{minipage}")
                            item_parts.append(r"\hfill")
                            item_parts.append(r"\begin{minipage}[t]{0.5\textwidth}")
                            item_parts.append(r"\vspace{8em}")  # 留白区域
                            item_parts.append(r"\end{minipage}")
                        elif diagram_content:
                            # 有答案时，图片正常显示
                            item_parts.append(self._wrap_diagram_block(diagram_content))
                        elif not include_answer:
                            # 没有图但需要留白
                            item_parts.append("\n" + r"\vspace{6em}")
                    else:
                        # 选填题：图在题干右侧（右对齐悬挂）
                        if diagram_content:
                            item_parts.append(self._wrap_diagram_block(diagram_content))
                    
                    # 答案和解析
                    if include_answer and q.answer:
                        item_parts.append(f"\n\\textbf{{答案：}} {self._escape_latex(q.answer)}")
                    if include_explanation and q.explanation:
                        item_parts.append(f"\n\\textbf{{解析：}} {self._escape_latex(q.explanation)}")
                    
                    section_content.append("\n".join(item_parts))
                    
                except Exception as e:
                    section_content.append(r"\item % 题目生成出错: " + str(e)[:50])
            
            section_content.append(r"\end{enumerate}")
            body_parts.append("\n".join(section_content))

        footer = r"""
\end{document}
"""
        return header + "\n\n".join(body_parts) + footer, attachments

    def compile_pdf(self, latex_content: str, attachments: List[Tuple[str, bytes]] | None = None) -> tuple[bool, str | Path, str]:
        """
        调用 xelatex 编译，返回 (ok, path/message, log_text)
        """
        import shutil

        engine = shutil.which("xelatex")
        if not engine:
            return False, "xelatex not found (please install texlive-xetex)", ""

        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                tmp_path = Path(tmpdir)
                tex_file = tmp_path / "paper.tex"
                tex_file.write_text(latex_content, encoding="utf-8")
                # 写入附件（图片等）
                for fname, data in attachments or []:
                    (tmp_path / fname).write_bytes(data)
                cmd = [
                    engine,
                    "-interaction=nonstopmode",
                    "-halt-on-error",
                    tex_file.name,
                ]
                proc = subprocess.run(
                    cmd,
                    cwd=tmp_path,
                    capture_output=True,
                    text=True,
                )
                log = proc.stdout + "\n" + proc.stderr
                pdf_file = tmp_path / "paper.pdf"
                # 只要 PDF 存在就算成功（LaTeX 警告会导致非零返回码）
                if pdf_file.exists():
                    # 安全创建临时文件保存结果
                    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp_out:
                        tmp_out.write(pdf_file.read_bytes())
                        out_file = Path(tmp_out.name)
                    return True, out_file, log
                return False, log or "xelatex failed", log
        except FileNotFoundError as exc:
            return False, "xelatex not found in PATH", str(exc)
        except Exception as exc:  # pragma: no cover - unexpected
            return False, f"compile error: {exc}", str(exc)

    def build_docx(
        self,
        paper: orm.Paper,
        pq_list: List[orm.PaperQuestion],
        question_map: Mapping[str, orm.Question],
        include_answer: bool = True,
        include_explanation: bool = True,
    ) -> tuple[bool, str | Path, str]:
        if not Document:
            return False, "python-docx not installed", ""
        try:
            doc = Document()
            doc.add_heading(paper.title, level=1)
            if paper.description:
                doc.add_paragraph(paper.description)

            for pq in sorted(pq_list, key=lambda x: x.order):
                q = question_map.get(pq.question_id)
                if not q:
                    continue
                p = doc.add_paragraph()
                p.add_run(f"{pq.order}. ({pq.score}分) ").bold = True
                p.add_run(q.question_text)
                if q.options:
                    for opt in q.options:
                        doc.add_paragraph(opt, style="List Bullet")
                if include_answer and q.answer:
                    doc.add_paragraph(f"【答案】{q.answer}")
                if include_explanation and q.explanation:
                    doc.add_paragraph(f"【解析】{q.explanation}")
                if q.has_geometry and q.geometry_svg:
                    svg_result = self._svg_to_png_attachment(q.geometry_svg)
                    if svg_result:
                        fname, data = svg_result
                        with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as img_tmp:
                            img_tmp.write(data)
                            image_path = Path(img_tmp.name)
                        doc.add_picture(image_path, width=None)
                        image_path.unlink(missing_ok=True)
                    else:
                        doc.add_paragraph("[SVG 未内嵌，请前端或后续步骤转换插入]")

            with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_out:
                doc.save(tmp_out.name)
                out_file = Path(tmp_out.name)
            return True, out_file, ""
        except Exception as exc:  # pragma: no cover
            return False, f"docx export error: {exc}", str(exc)

    def cleanup_file(self, path: str | Path):
        try:
            Path(path).unlink(missing_ok=True)
        except Exception:
            pass

    def _svg_to_png_attachment(self, svg_content: str) -> tuple[str, bytes] | None:
        """
        尝试将 SVG 转 PNG，失败则返回 None。
        """
        if not svg_content:
            return None
        if cairosvg is None:  # pragma: no cover - optional dependency not installed
            return None
        try:
            png_bytes = cairosvg.svg2png(bytestring=svg_content.encode("utf-8"))
            fname = f"svg_{uuid.uuid4().hex}.png"
            return fname, png_bytes
        except Exception:
            return None

    def _svg_to_pdf_attachment(self, svg_content: str) -> tuple[str, bytes] | None:
        """
        将 SVG 转 PDF（矢量格式），用于 LaTeX includegraphics。
        相比 PNG：文件更小，矢量缩放不失真。
        """
        if not svg_content:
            return None
        if cairosvg is None:  # pragma: no cover - optional dependency not installed
            return None
        try:
            pdf_bytes = cairosvg.svg2pdf(bytestring=svg_content.encode("utf-8"))
            fname = f"svg_{uuid.uuid4().hex}.pdf"
            return fname, pdf_bytes
        except Exception:
            return None

    def _svg_to_tikz_block(self, svg_content: str) -> str | None:
        """
        将简单 SVG 转换为 TikZ 片段（支持 line/circle/ellipse/text 基础元素）。
        若无法解析则返回 None。
        """
        if not svg_content:
            return None
        try:
            root = ET.fromstring(svg_content)
        except Exception:
            return None

        # 获取画布大小用于翻转 y 轴
        width, height = 400.0, 400.0
        viewbox = root.get("viewBox")
        if viewbox:
            parts = viewbox.replace(",", " ").split()
            if len(parts) == 4:
                try:
                    width = float(parts[2])
                    height = float(parts[3])
                except Exception:
                    pass
        else:
            try:
                width = float((root.get("width") or "400").replace("px", ""))
                height = float((root.get("height") or "400").replace("px", ""))
            except Exception:
                pass

        # 收集 defs 内的元素（通常是箭头 marker 等装饰），转换时跳过
        defs_nodes = set()
        for defs in root.findall(".//{*}defs"):
            defs_nodes.update(list(defs.iter()))

        cmds: List[str] = []
        scale = 0.03  # 将 400x400 缩放到约 12x12

        def fmt(coord: str | None) -> float:
            try:
                return float(coord or "0.0")
            except Exception:
                return 0.0

        def flip_y(y: float) -> float:
            return (height - y) * scale

        def is_dashed(el: ET.Element) -> bool:
            style = el.get("style", "")
            cls = el.get("class", "")
            dasharray = el.get("stroke-dasharray", "")
            return ("dash" in style) or ("dash" in cls) or (dasharray not in ("", None))

        def parse_path(d: str) -> List[List[tuple[float, float]]]:
            """
            粗略解析 path 数据，支持直线/水平/垂直/二次、三次贝塞尔曲线。
            返回若干折线段，每段为点列表，后续用 -- 连接。
            复杂的弧线（A/a）会退化为直线连接起终点。
            """
            import re

            def tokenize(data: str) -> List[str]:
                return re.findall(r"[MmLlHhVvCcSsQqTtAaZz]|-?\d*\.?\d+(?:[eE][-+]?\d+)?", data or "")

            def is_cmd(tok: str) -> bool:
                return len(tok) == 1 and tok.isalpha()

            def read_numbers(n: int) -> List[float]:
                nonlocal idx
                vals = []
                for _ in range(n):
                    if idx >= len(tokens):
                        break
                    vals.append(float(tokens[idx]))
                    idx += 1
                return vals

            def cubic_samples(p0, p1, p2, p3, steps: int = 10):
                out = []
                for i in range(1, steps + 1):
                    t = i / steps
                    mt = 1 - t
                    x = (
                        mt * mt * mt * p0[0]
                        + 3 * mt * mt * t * p1[0]
                        + 3 * mt * t * t * p2[0]
                        + t * t * t * p3[0]
                    )
                    y = (
                        mt * mt * mt * p0[1]
                        + 3 * mt * mt * t * p1[1]
                        + 3 * mt * t * t * p2[1]
                        + t * t * t * p3[1]
                    )
                    out.append((x, y))
                return out

            def quad_samples(p0, p1, p2, steps: int = 10):
                out = []
                for i in range(1, steps + 1):
                    t = i / steps
                    mt = 1 - t
                    x = mt * mt * p0[0] + 2 * mt * t * p1[0] + t * t * p2[0]
                    y = mt * mt * p0[1] + 2 * mt * t * p1[1] + t * t * p2[1]
                    out.append((x, y))
                return out

            tokens = tokenize(d)
            segments: List[List[tuple[float, float]]] = []
            current: List[tuple[float, float]] = []
            idx = 0
            cmd = ""
            cursor = (0.0, 0.0)
            start_point = (0.0, 0.0)
            last_ctrl: tuple[float, float] | None = None

            def move_to(pt: tuple[float, float]):
                nonlocal cursor, start_point, current
                if current:
                    segments.append(current)
                current = [pt]
                cursor = pt
                start_point = pt

            def line_to(pt: tuple[float, float]):
                nonlocal cursor
                current.append(pt)
                cursor = pt

            while idx < len(tokens):
                if is_cmd(tokens[idx]):
                    cmd = tokens[idx]
                    idx += 1
                if cmd == "":
                    break

                abs_cmd = cmd.upper()
                is_relative = cmd.islower()

                if abs_cmd == "M":
                    nums = read_numbers(2)
                    if len(nums) < 2:
                        break
                    x, y = nums
                    if is_relative:
                        x += cursor[0]
                        y += cursor[1]
                    move_to((x, y))
                    # 额外的坐标对视为连续 lineto
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        extra = read_numbers(2)
                        if len(extra) < 2:
                            break
                        ex, ey = extra
                        if is_relative:
                            ex += cursor[0]
                            ey += cursor[1]
                        line_to((ex, ey))
                    last_ctrl = None
                    cmd = "L" if abs_cmd == "M" else "l"
                elif abs_cmd == "L":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(2)
                        if len(nums) < 2:
                            break
                        x, y = nums
                        if is_relative:
                            x += cursor[0]
                            y += cursor[1]
                        line_to((x, y))
                    last_ctrl = None
                elif abs_cmd == "H":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(1)
                        if not nums:
                            break
                        x = nums[0]
                        x = x + cursor[0] if is_relative else x
                        line_to((x, cursor[1]))
                    last_ctrl = None
                elif abs_cmd == "V":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(1)
                        if not nums:
                            break
                        y = nums[0]
                        y = y + cursor[1] if is_relative else y
                        line_to((cursor[0], y))
                    last_ctrl = None
                elif abs_cmd == "C":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(6)
                        if len(nums) < 6:
                            break
                        x1, y1, x2, y2, x, y = nums
                        if is_relative:
                            x1 += cursor[0]; y1 += cursor[1]
                            x2 += cursor[0]; y2 += cursor[1]
                            x += cursor[0]; y += cursor[1]
                        samples = cubic_samples(cursor, (x1, y1), (x2, y2), (x, y))
                        for pt in samples:
                            line_to(pt)
                        cursor = (x, y)
                        last_ctrl = (x2, y2)
                elif abs_cmd == "S":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(4)
                        if len(nums) < 4:
                            break
                        x2, y2, x, y = nums
                        if is_relative:
                            x2 += cursor[0]; y2 += cursor[1]
                            x += cursor[0]; y += cursor[1]
                        if last_ctrl is None:
                            x1, y1 = cursor
                        else:
                            x1 = 2 * cursor[0] - last_ctrl[0]
                            y1 = 2 * cursor[1] - last_ctrl[1]
                        samples = cubic_samples(cursor, (x1, y1), (x2, y2), (x, y))
                        for pt in samples:
                            line_to(pt)
                        cursor = (x, y)
                        last_ctrl = (x2, y2)
                elif abs_cmd == "Q":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(4)
                        if len(nums) < 4:
                            break
                        x1, y1, x, y = nums
                        if is_relative:
                            x1 += cursor[0]; y1 += cursor[1]
                            x += cursor[0]; y += cursor[1]
                        samples = quad_samples(cursor, (x1, y1), (x, y))
                        for pt in samples:
                            line_to(pt)
                        cursor = (x, y)
                        last_ctrl = (x1, y1)
                elif abs_cmd == "T":
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(2)
                        if len(nums) < 2:
                            break
                        x, y = nums
                        if is_relative:
                            x += cursor[0]; y += cursor[1]
                        if last_ctrl is None:
                            x1, y1 = cursor
                        else:
                            x1 = 2 * cursor[0] - last_ctrl[0]
                            y1 = 2 * cursor[1] - last_ctrl[1]
                        samples = quad_samples(cursor, (x1, y1), (x, y))
                        for pt in samples:
                            line_to(pt)
                        cursor = (x, y)
                        last_ctrl = (x1, y1)
                elif abs_cmd == "A":
                    # 复杂弧线退化为直线到终点
                    while idx < len(tokens) and not is_cmd(tokens[idx]):
                        nums = read_numbers(7)
                        if len(nums) < 7:
                            break
                        # rx, ry, x-axis-rotation, large-arc-flag, sweep-flag, x, y
                        x = nums[-2]
                        y = nums[-1]
                        if is_relative:
                            x += cursor[0]; y += cursor[1]
                        line_to((x, y))
                        cursor = (x, y)
                        last_ctrl = None
                elif abs_cmd == "Z":
                    # 闭合当前子路径
                    if current and (current[-1] != start_point):
                        current.append(start_point)
                    if current:
                        segments.append(current)
                        current = []
                    cursor = start_point
                    last_ctrl = None
                else:
                    # 未支持的命令，跳过
                    idx += 1
                    last_ctrl = None

            if current:
                segments.append(current)
            return segments

        for el in root.iter():
            if el in defs_nodes:
                continue  # 跳过 marker/defs 内部的装饰元素
            tag = el.tag.split("}")[-1].lower()
            dashed = "[dashed]" if is_dashed(el) else ""
            if tag == "line":
                x1, y1 = fmt(el.get("x1")), fmt(el.get("y1"))
                x2, y2 = fmt(el.get("x2")), fmt(el.get("y2"))
                cmds.append(r"\draw%s (%.3f,%.3f) -- (%.3f,%.3f);" % (dashed, x1 * scale, flip_y(y1), x2 * scale, flip_y(y2)))
            elif tag == "circle":
                cx, cy = fmt(el.get("cx")), fmt(el.get("cy"))
                r = fmt(el.get("r"))
                cmds.append(r"\draw%s (%.3f,%.3f) circle (%.3f);" % (dashed, cx * scale, flip_y(cy), r * scale))
            elif tag == "ellipse":
                cx, cy = fmt(el.get("cx")), fmt(el.get("cy"))
                rx, ry = fmt(el.get("rx")), fmt(el.get("ry"))
                cmds.append(r"\draw%s (%.3f,%.3f) ellipse (%.3f and %.3f);" % (dashed, cx * scale, flip_y(cy), rx * scale, ry * scale))
            elif tag == "path":
                segments = parse_path(el.get("d") or "")
                for seg in segments:
                    if len(seg) >= 2:
                        coords = " -- ".join(["(%.3f,%.3f)" % (x * scale, flip_y(y)) for x, y in seg])
                        cmds.append(r"\draw%s %s;" % (dashed, coords))
            elif tag == "rect":
                # 矩形：左上角 (x, y)，宽 width，高 height
                x, y = fmt(el.get("x")), fmt(el.get("y"))
                w, h = fmt(el.get("width")), fmt(el.get("height"))
                # 只绘制有 stroke 的矩形，忽略纯背景
                stroke = el.get("stroke", "")
                if stroke and stroke.lower() != "none":
                    x1, y1 = x * scale, flip_y(y)
                    x2, y2 = (x + w) * scale, flip_y(y + h)
                    cmds.append(r"\draw%s (%.3f,%.3f) rectangle (%.3f,%.3f);" % (dashed, x1, y1, x2, y2))
            elif tag == "polygon":
                # 多边形：points="x1,y1 x2,y2 x3,y3 ..."
                points_str = el.get("points", "")
                pts = []
                for pt in points_str.replace(",", " ").split():
                    try:
                        pts.append(float(pt))
                    except:
                        continue
                if len(pts) >= 6:  # 至少 3 个点
                    coords = []
                    for i in range(0, len(pts), 2):
                        if i + 1 < len(pts):
                            coords.append("(%.3f,%.3f)" % (pts[i] * scale, flip_y(pts[i+1])))
                    if coords:
                        cmds.append(r"\draw%s %s -- cycle;" % (dashed, " -- ".join(coords)))
            elif tag == "text":
                x, y = fmt(el.get("x")), fmt(el.get("y"))
                dx = fmt(el.get("dx"))
                dy = -fmt(el.get("dy"))  # dy 需要翻转
                txt = self._normalize_math_content((el.text or "").strip())
                if txt:
                    # 用数学模式包裹文字（适合坐标点标签）
                    cmds.append(r"\node at (%.3f,%.3f) {$%s$};" % ((x + dx) * scale, flip_y(y) + dy * scale, txt))

        if not cmds:
            return None
        # 高考卷风格：居中、适当缩放、细线条
        tikz = ["\\begin{tikzpicture}[>=Stealth, scale=0.8, line width=0.5pt]", *cmds, "\\end{tikzpicture}"]
        return "\n".join(tikz)

    def svg_to_png_base64(self, svg_content: str) -> str | None:
        """
        将 SVG 转 base64 PNG，前端可直接展示。
        """
        if not svg_content or cairosvg is None:
            return None
        try:
            png_bytes = cairosvg.svg2png(bytestring=svg_content.encode("utf-8"))
            return f"data:image/png;base64,{base64.b64encode(png_bytes).decode('utf-8')}"
        except Exception:
            return None

    def _wrap_diagram_block(self, content: str) -> str:
        """
        将图形包裹在 minipage 中，默认居右，避免占满版心。
        """
        if not content:
            return ""
        return (
            "\n\\par\\noindent\\hfill\\begin{minipage}{0.45\\textwidth}\\centering\n"
            + content
            + "\n\\end{minipage}\\hfill\\null\n"
        )

    def _clean_markdown(self, text: str) -> str:
        """
        简单去掉常见的 Markdown 标记，保留公式/纯文本。
        同时统一全角标点为半角（保持括号样式一致）。
        """
        if not text:
            return ""
        t = text
        # 去掉代码块
        t = re.sub(r"```.*?```", "", t, flags=re.S)
        # 去掉标题符号
        t = re.sub(r"^\\s*#{1,6}\\s*", "", t, flags=re.M)
        # 去掉加粗/斜体标记
        t = t.replace("**", "") # 保留 __ 以防误删填空下划线
        # 去掉列表标记
        t = re.sub(r"^\\s*[-+*]\\s+", "", t, flags=re.M)
        
        # 统一全角标点为半角（保持括号/逗号/冒号等样式一致）
        fullwidth_to_halfwidth = {
            "（": "(",
            "）": ")",
            "，": ", ",  # 全角逗号转半角并加空格
            "：": ": ",  # 全角冒号转半角并加空格  
            "；": "; ",  # 全角分号转半角并加空格
            "？": "?",
            "！": "!",
            "．": ".",
            "　": " ",   # 全角空格转半角
        }
        for fw, hw in fullwidth_to_halfwidth.items():
            t = t.replace(fw, hw)
        
        return t.strip()

    def _escape_latex(self, text: str) -> str:
        """
        转义特殊字符，但保留数学环境 $...$ 和 $$...$$ 内的内容不转义。
        自动检测并修复未闭合的 $ 符号。
        """
        if not text:
            return ""
        
        # 先简单清洗 Markdown
        text = self._clean_markdown(text)

        # 自动修复未闭合的 $ 符号
        # 计算 $ 数量（排除转义的 \$）
        dollar_count = len(re.findall(r'(?<!\\)\$', text))
        if dollar_count % 2 != 0:
            # 奇数个 $，说明有未闭合的，在末尾补一个
            text = text + '$'
        
        # 使用正则分割，保留数学环境
        # 匹配 $$...$$ 或 $...$（非贪婪）
        pattern = r'(\$\$.*?\$\$|\$.*?\$)'
        parts = re.split(pattern, text, flags=re.DOTALL)
        
        def _normalize_plain(t: str) -> str:
            # 常见符号/习惯用法替换为 LaTeX 形式（仅在文本环境处理）
            replacements = {
                "π": r"$\pi$",
                "∥": r"$\spar$",
                "∞": r"$\infty$",
                "×": r"$\times$",
                "÷": r"$\div$",
                "°": r"$^\circ$",
            }
            for k, v in replacements.items():
                t = t.replace(k, v)
            t = re.sub(r"\s//\s", r" \\spar ", t)
            return t
        
        result = []
        for i, part in enumerate(parts):
            if part.startswith('$$') or part.startswith('$'):
                # 数学环境，直接保留
                result.append(self._normalize_math_content(part))
            else:
                # 非数学环境，转义特殊字符
                normalized = _normalize_plain(part)
                escaped = self._escape_text_only(normalized)
                result.append(escaped)
        
        return ''.join(result)
    
    def _escape_text_only(self, text: str) -> str:
        """
        仅转义普通文本中的特殊字符（不在数学环境中）。
        连续下划线（____）作为填空横线处理。
        """
        # 使用占位符保护填空横线（多种格式）
        BLANK_PLACEHOLDER = "\x00BLANK\x00"
        
        # 处理各种填空横线格式：
        # 1. 连续下划线 ____ 
        # 2. 已转义的 \_\_\_\_ 
        # 3. 混合格式
        text = re.sub(r'(\\?_){2,}', BLANK_PLACEHOLDER, text)
        
        # 转义特殊字符
        replacements = {
            "&": r"\&",
            "%": r"\%",
            "#": r"\#",
            "_": r"\_",
            "{": r"\{",
            "}": r"\}",
            "~": r"\textasciitilde{}",
            "^": r"\^{}",
        }
        out = []
        for ch in text:
            out.append(replacements.get(ch, ch))
        result = "".join(out)
        
        # 将占位符替换为 LaTeX 填空横线命令
        result = result.replace(BLANK_PLACEHOLDER, r"\undsp ")
        return result

    def _strip_option_prefix(self, text: str) -> str:
        """
        去掉选项开头的 A./A．/A、 等前缀，避免与模板宏重复。
        """
        if not text:
            return ""
        return re.sub(r"^\s*[A-DＡ-Ｄa-d][\.\。．、﹒\)]\s*", "", text).strip()

    def _normalize_math_content(self, text: str) -> str:
        """
        将常见的 Unicode 符号替换为 LaTeX 数学命令，供数学环境或 TikZ 节点使用。
        """
        if not text:
            return ""
        # Unicode 符号替换
        replacements = {
            "π": r"\pi",
            "∥": r"\spar",
            "∞": r"\infty",
            "×": r"\times",
            "÷": r"\div",
            "°": r"^\circ",
            # Unicode 下标数字
            "₀": "_0", "₁": "_1", "₂": "_2", "₃": "_3", "₄": "_4",
            "₅": "_5", "₆": "_6", "₇": "_7", "₈": "_8", "₉": "_9",
            # Unicode 上标数字
            "⁰": "^0", "¹": "^1", "²": "^2", "³": "^3", "⁴": "^4",
            "⁵": "^5", "⁶": "^6", "⁷": "^7", "⁸": "^8", "⁹": "^9",
        }
        for k, v in replacements.items():
            text = text.replace(k, v)
        # 三角函数转为 LaTeX 命令（仅替换独立的函数名，避免误替换变量名）
        # 使用 negative lookbehind (?<!\\) 确保前面没有反斜杠，避免 \sin 变成 \\sin
        trig_funcs = ["sin", "cos", "tan", "cot", "sec", "csc", "arcsin", "arccos", "arctan", "ln", "log", "exp"]
        for fn in trig_funcs:
            # 匹配单词边界，且前面不能有反斜杠（避免重复转义）
            text = re.sub(rf"(?<!\\)\b{fn}\b", rf"\\{fn}", text)
        # 将惯用的 // 视为平行符号，避免 URL 误替换：排除前面有冒号或反斜杠的情况
        text = re.sub(r"(?<!:)(?<!\\)//", r"\\spar ", text)
        return text
